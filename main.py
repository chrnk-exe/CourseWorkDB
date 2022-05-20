from sqlalchemy import create_engine
from sqlalchemy import Column, Integer, String, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from docx import *
import time
import json


categories = {'A': 100, 'B': 200, 'C': 300}
categories_price = {'A': 3, 'B': 2, 'C': 1}
secret_key = 'admin123'
engine = create_engine('sqlite:///db.db', echo=False)
Session = sessionmaker(bind=engine)
session = Session()
Base = declarative_base()


class Cinema(Base):
    __tablename__ = 'Cinema'
    id = Column(Integer, primary_key=True)
    name = Column(String)
    district = Column(String)
    address = Column(String)
    category = Column(String)
    capacity = Column(String)
    sessions = Column(Text, default='[]')

    def __init__(self, name, district, address, category):
        self.name = name
        self.district = district
        self.address = address
        self.category = category
        self.capacity = categories[category]


class Film(Base):
    __tablename__ = 'Films'
    id = Column(Integer, primary_key=True)
    name = Column(String)
    country = Column(String)
    director = Column(String)
    genre = Column(String)
    price = Column(Integer)

    def __init__(self, name, country, director, genre, price):
        self.name = name
        self.country = country
        self.director = director
        self.genre = genre
        self.price = price


class Seance(Base):
    __tablename__ = 'Sessions'
    id = Column(Integer, primary_key=True)
    cinema = Column(String)
    time = Column(String)
    movie_name = Column(String)
    tickets = Column(Integer)

    def __init__(self, cinema, time, movie_name, tickets):
        self.cinema = cinema
        self.time = time
        self.movie_name = movie_name
        self.tickets = tickets


# Создание таблицы
Base.metadata.create_all(engine)


def add_cinema(name, district, address, category):
    cinema = Cinema(name, district, address, category)
    session.add(cinema)
    session.commit()


def add_film(name, country, director, genre, price):
    film = Film(name, country, director, genre, price)
    session.add(film)
    session.commit()


def add_session(cinema_name, time, movie_name):

    cinema = session.query(Cinema).filter_by(name=cinema_name).first()

    sess = Seance(cinema_name, time, movie_name, cinema.capacity)
    session.add(sess)
    session.commit()

    sess_list = json.loads(cinema.sessions)
    sess_list.append(sess.id)
    cinema.sessions = json.dumps(sess_list)
    session.commit()


def delete_cinema(cinema_name):
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        session_list = json.loads(cinema.sessions)
        session.delete(cinema)
    except Exception as ex:
        print('Такого кинотеатра не существует')
        return
    if len(session_list) != 0:
        for index in session_list:
            sess = session.query(Seance).filter_by(id=index).first()
            session.delete(sess)
    session.commit()
    print('Удаление завершено успешно')


def get_tickets_count():
    cinema_name = input('Введите название кинотеатра: ')
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        session_list = json.loads(cinema.sessions)
        for index in session_list:
            sess = session.query(Seance).filter_by(id=index).first()
            print(str(sess.id) + '. фильм: ' + str(sess.movie_name) + ' количество билетов: ' + str(sess.tickets) + ' дата сеанса: ' + str(sess.time))
        return sess.tickets
    except Exception as ex:
        print('Такого кинотеатра не существует')
        return 0


def buy_ticket():
    cinema_name = input('Введите назавние кинотеатра: ')
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    session_list = json.loads(cinema.sessions)
    for session_index in session_list:
        se = session.query(Seance).filter_by(id=session_index).first()
        print(str(se.id) + '. ' + 'Дата сеанса - ' + str(se.time) + ' Количество билетов - ' + str(se.tickets))
    session_index = input('Введите номер сеанса: ')
    sess = session.query(Seance).filter_by(id=session_index).first()
    if sess.tickets == 0:
        print('Билеты все куплены, выберите другой сеанс!')
        return 0
    sess.tickets -= 1
    session.commit()
    return 0


def get_ticket_price():
    cinema_name = input('Введите название кинотеатра: ')
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        session_list = json.loads(cinema.sessions)
        for index in session_list:
            sess = session.query(Seance).filter_by(id=index).first()
            print(str(sess.id) + '. фильм: ' + str(sess.movie_name) + ' количество билетов: ' + str(sess.tickets) + ' дата сеанса: ' + str(sess.time))
        seance_index = input('Введите индекс сеанса: ')
        sess = session.query(Seance).filter_by(id=seance_index).first()
        film = session.query(Film).filter_by(name=sess.movie_name).first()
        print('Цена билета: ' + str(film.price * categories_price[cinema.category]))
    except Exception as ex:
        return 0


def get_film_info(film_name):
    film = session.query(Film).filter_by(name=film_name).first()
    try:
        return [film.genre, film.country, film.director]
    except Exception as ex:
        return 0


def cinema_capacity(cinema_name):
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        return cinema.capacity
    except Exception as ex:
        return 0


def remove_film(name):
    film = session.query(Film).filter_by(name=name).first()
    try:
        session.delete(film)
    except Exception as ex:
        return 666
    session.commit()
    return 0


def change_category(name, category):
    cinema = session.query(Cinema).filter_by(name=name).first()
    cinema.category = category
    session.commit()


def get_cinema_repertoire(cinema_name):
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        session_list = json.loads(cinema.sessions)
    except Exception as ex:
        return 0

    response = set()
    for session_index in session_list:
        sess = session.query(Seance).filter_by(id=session_index).first()
        response.add(sess.movie_name)

    return response


def get_address_and_district(cinema_name):
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        return [cinema.address, cinema.district]
    except Exception as ex:
        return 0


def get_film_sessions(cinema_name, film_name):
    cinema = session.query(Cinema).filter_by(name=cinema_name).first()
    try:
        session_list = json.loads(cinema.sessions)
    except Exception as ex:
        return 0

    response = []
    for session_index in session_list:
        sess = session.query(Seance).filter_by(id=session_index).first()
        if sess.movie_name == film_name and sess.id not in response:
            # list(sess[index] for index in sess)
            response.append([sess.cinema, sess.time, sess.movie_name, sess.tickets])

    return response


def get_district_report():

    districts = session.query(Cinema.district).all()
    districts = set(districts)
    districts_films_info = []
    for district in districts:
        films_info = {}
        cinemas = session.query(Cinema).filter_by(district=district[0]).all()
        for cinema in cinemas:

            session_list = json.loads(cinema.sessions)
            for session_index in session_list:
                sess = session.query(Seance).filter_by(id=session_index).first()
                film = session.query(Film).filter_by(name=sess.movie_name).first()
                price = film.price * categories_price[cinema.category]

                if sess.movie_name in films_info:
                    if cinema.name not in films_info[sess.movie_name]:
                        films_info[sess.movie_name].append({cinema.name: price})
                else:
                    films_info[sess.movie_name] = [{cinema.name: price}]
        districts_films_info.append({district[0]: films_info})

    return districts_films_info


def show_menu(condition):
    if condition:
        print('1. Получить репертуар кинотеатра (по названию кинотеатра)')
        print('2. Получить адрес и район кинотеатра (по названию кинотеатра); ')
        print('3. число мест (свободных) на данный сеанс (название кинотеатра и сеанс)')
        print('4. цена билетов на данный сеанс (название кинотеатра и сеанс);')
        print('5. жанр, производство и режиссер данного фильма (по названию); ')
        print('6. вместимость данного кинотеатра (по названию кинотеатра). ')
        print('Изменения: ')
        print('7. открытие нового кинотеатра')
        print('8. снятие фильма с проката')
        print('9. добавить фильм в репертуар')
        print('10. Добавить сеанс кинотеатру')
        print('11. Купить билет на сеанс')
        print('12. Удалить информацио о кинотеатре')
        print('Отчёты: ')
        print('13. Сеансы в кинотеатре (по кинотеатру и фильму)')
        print('14. отчет о прокате фильмов в районах города')
    else:
        print('1. Получить репертуар кинотеатра (по названию кинотеатра)')
        print('2. Получить адрес и район кинотеатра (по названию кинотеатра); ')
        print('3. число мест (свободных) на данный сеанс (название кинотеатра и сеанс)')
        print('4. цена билетов на данный сеанс (название кинотеатра и сеанс);')
        print('5. жанр, производство и режиссер данного фильма (по названию); ')
        print('6. вместимость данного кинотеатра (по названию кинотеатра). ')
        print('Отчёты: ')
        print('13. Сеансы в кинотеатре (по кинотеатру и фильму)')
        print('14. отчет о прокате фильмов в районах города')


def create_film_session_report(film_name, cinema_name):
    report = get_film_sessions(cinema_name, film_name)
    if report:
        document = Document()
        document.add_heading('Отчёт о сеансах ' + str(film_name) + ' в кинотеатре ' + str(cinema_name), 0)
        table = document.add_table(rows=1, cols=4)
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Кинотеатр'
        header_cells[1].text = 'Время'
        header_cells[2].text = 'Фильм'
        header_cells[3].text = 'Количество мест на сеанс'
        for seance in report:
            row_cells = table.add_row().cells
            for i in range(len(row_cells)):
                row_cells[i].text = str(seance[i])
        l = time.asctime().split(' ')
        l = '-'.join(l)
        l = l.split(':')
        l = '_'.join(l)
        document.save('film-session-report-'+l+'.docx')
    else:
        return 0


def create_districts_report():
    report = get_district_report()
    document = Document()
    l = time.asctime().split(' ')
    l = '-'.join(l)
    l = l.split(':')
    l = '_'.join(l)
    document.add_heading('Отчёт о прокате фильмов ', 0)
    table = document.add_table(rows=1, cols=4)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Районы'
    header_cells[1].text = 'Фильмы'
    header_cells[2].text = 'Кинотеатр'
    header_cells[3].text = 'Цена билета'
    for dist in report:
        row_cells = table.add_row().cells
        row_cells[0].text = list(dist.keys())[0]
        for films in dist.values():
            for film, cinema_price in films.items():
                rw_cells = table.add_row().cells
                rw_cells[1].text = film
                for cinemas in cinema_price:
                    for cinema, price in cinemas.items():
                        r_cells = table.add_row().cells
                        r_cells[2].text = cinema
                        r_cells[3].text = str(price)
    document.save('district-report-' + l + '.docx')
    return 0


def main():
    print('1. Сотрудник справочной службы')
    print('2. Администратор БД')
    secret = int(input('Выберите пользователя (1 или 2): '))
    password = ''
    if secret == 2:
        password = input('Введите секретный ключ: ')
    while password == secret_key:
        show_menu(True)
        try:
            act = int(input('Введите номер нужного действия (для выхода напишите 0): '))
        except ValueError:
            print('Непреднамеренный выход из программы')
            return
        if act == 1:
            name = input('Введите имя кинотеатра: ')
            result = get_cinema_repertoire(name)
            if result:
                for film in result:
                    print(film, end='')
            else:
                print('Такого кинотеатра не существует')
        elif act == 2:
            name = input('Введите имя кинотеатра: ')
            result = get_address_and_district(name)
            if result:
                for i in result:
                    print(i, end=' ')
            else:
                print('Такого кинотеатра не существует')
            print()
        elif act == 3:
            get_tickets_count()
        elif act == 4:
            get_ticket_price()
        elif act == 5:
            name = input('Введите название фильма: ')
            result = get_film_info(name)
            if result:
                print('Жанр - ' + result[0])
                print('Производство - ' + result[1])
                print('Режиссёр - ' + result[2])
            else:
                print('такого фильма нет')
        elif act == 6:
            name = input('Введите название кинотеатра: ')
            result = cinema_capacity(name)
            if result:
                print(result)
            else:
                print('Такого кинотеатра НЕ существует')
        elif act == 7:
            print('Введите информацию о кинотеатре')
            name = input('Название: ')
            district = input('Район: ')
            address = input('Адрес: ')
            category = input('Категория (A, B или C): ')
            add_cinema(name, district, address, category)
        elif act == 8:
            name = input('Введите название фильма: ')
            result = remove_film(name)
            if result:
                print('Данного фильма уже не существует в репертуаре кинотеатров')
        elif act == 9:
            print('Введите информацию о фильме')
            name = input('Название: ')
            country = input('Производство: ')
            director = input('Режиссёр: ')
            genre = input('Жанр: ')
            price = input('Цена билетов: ')
            add_film(name, country, director, genre, price)
        elif act == 10:
            cinema_name = input('Название кинотеатра: ')
            seance_year = input('Введите год: ')
            seance_month = input('Введите месяц сеанса (1-12): ')
            seance_day = input('Введите день сеанса (1-31): ')
            seance_time = input('Введите время (чч:мм): ')
            film_name = input('Название фильма: ')
            seance_main_time = ' '.join([seance_year, seance_month, seance_day, seance_time])
            add_session(cinema_name, seance_main_time, film_name)
        elif act == 11:
            buy_ticket()
        elif act == 12:
            print('!ВНИМАНИЕ!')
            print('При удалении кинотеатра из базы данных удаляются и все связанные с ним сеансы')
            cinema_name = input('Введите название кинотеатра: ')
            delete_cinema(cinema_name)
        elif act == 13:

            name_film = input('назвние фильма: ')
            cinema_name = input('название кинотеатра: ')
            create_film_session_report(name_film, cinema_name)
            result = get_film_sessions(cinema_name, name_film)
            if result:
                print(result)
            else:
                print('Такого кинотеатра НЕ существует')
        elif act == 14:
            create_districts_report()

        try:
            a = input('Для продолжения введите любой символ: ')
        except ValueError:
            print('Выход из программы')
        if not act:
            break
    while secret == 1:
        show_menu(False)
        try:
            act = int(input('Введите номер нужного действия (для выхода напишите 0): '))
        except ValueError:
            print('Непреднамеренный выход из программы')
            return
        if act == 1:
            name = input('Введите имя кинотеатра: ')
            result = get_cinema_repertoire(name)
            if result:
                for film in result:
                    print(film, end=', ')
            else:
                print('Такого кинотеатра не существует')
            print()
        elif act == 2:
            name = input('Введите имя кинотеатра: ')
            result = get_address_and_district(name)
            if result:
                for i in result:
                    print(i, end=' ')
            else:
                print('Такого кинотеатра не существует')
            print()
        elif act == 3:
            get_tickets_count()
        elif act == 4:
            get_ticket_price()
        elif act == 5:
            name = input('Введите название фильма: ')
            result = get_film_info(name)
            if result:
                print('Жанр - ' + result[0])
                print('Производство - ' + result[1])
                print('Режиссёр - ' + result[2])
            else:
                print('такого фильма нет')
        elif act == 6:
            name = input('Введите название кинотеатра: ')
            result = cinema_capacity(name)
            if result:
                print(result)
            else:
                print('Такого кинотеатра НЕ существует')
        elif act == 13:

            name_film = input('назвние фильма: ')
            cinema_name = input('название кинотеатра: ')
            create_film_session_report(name_film, cinema_name)
            result = get_film_sessions(cinema_name, name_film)
            if result:
                print(result)
            else:
                print('Такого кинотеатра НЕ существует')
        elif act == 14:
            create_districts_report()

if __name__ == '__main__':
    main()
